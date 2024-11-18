'''
Student ID: 20106991
Jeisson Nino

Assessmenet 3 part B

3/06/2024
'''

from docx import Document
from docx.shared import Inches
import os
import tkinter as tk
from tkinter import filedialog

'''
1.1.	Create a word document with at least two headings.
1.2.	Add paragraphs to headings from task 1.1.
1.3.	Add a picture to your word document.
1.4.	Get full text from word document.c

'''


#-------------------------- FUNCTIONS ------------------------

# This function will clear the current console window
def clear_console():
    os.system('cls' if os.name == 'nt' else 'clear')

# this function will open a file explorer to select a file from a directory
def select_file():
    try:
        root = tk.Tk()
        root.withdraw()  # Hides the small tkinter window
        file_path = filedialog.askopenfilename()  # Opens the file dialog prompt
        if not file_path:  # Check if the file selection was cancelled
            print("File selection cancelled.")
            return None
        return file_path
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

# this function will create a Word Document
def create_doc():

    #Create a new word document and name it
    document = Document()
    file_name = input('What will be the name of the file?: \n')
    

    #Adding heading
    heading1= input('What will be the heading 1: \n')
    document.add_heading(heading1, level=0)

    #Add Paragraph2
    p1 = input('Paragraph 1: \n')
    document.add_paragraph(p1)

    heading2 = input('What will be the heading 2: \n')
    document.add_heading(heading2, level=1)

    #Add Paragraph2
    p2 = input('Paragraph 2: \n')
    document.add_paragraph(p2)

    document.save(file_name+'.docx')

    #return file_name


# This function will add a picture to a Word Docx
def add_picture():

    print('----------------Menu Add picture--------------------')
    
    print('Open the browser to add a picture')
    input()
    #selecting the picture from the explorer
    picture_path=select_file()
    print(picture_path)
    if not picture_path:
        print('No picture selected existing')
        return

    #Selecting the document from the explorer
    print('Select the document to add the picture: \n')
    input()
    document_path = select_file()
    print(document_path)
    if not document_path:
        print('no document selected existing')
    
    #Selecting the document from the explorer

    try: 
        document = Document(document_path)
        document.add_picture(picture_path, width= Inches(1.5), height= Inches(1.5))
        document.save('testing_word.docx')
        input('Document saved, enter to continue: ')
    except Exception as e:
        print(f'Failed to add the picture or saving the document: {e}')
    


# This function will take the text from a Word Docx and extract it in binary
# to be printed in the console
def get_text_doc():

    input('Select the document you want to extract: ')
    docfileName = select_file()
    print(docfileName)

    document = Document(docfileName)
    finalText= []
    for line in document.paragraphs:
        finalText.append(line.text)

    print('The text to print is: ')
    print('-----------------------------------------------------------------------')
    return print('\n'.join(finalText))



def display_menu():
    
    print('##############################################')
    print('             Word Docxs Manager MENU           ')
    print('\n1. Create Word doc.')
    print('2. Add picture to doc.')
    print('3. Get text from the Doc.')
    print('4. Exit.')
    

#--------------------------------------------------------------------------
def user_choice():
    while True:
        try:
            choice = int(input('Enter your choice (1-4): '))
            clear_console()
            if  choice < 1 or choice > 5:
                print('Invalid input. Please try again (1-4).')
                raise ValueError('Invalid choice')
            else:
                return choice
        except ValueError:
            print('Invalid input. Please try again (1-4).')


#--------------------------------------------------------------------------
def menu():
    while True:
        #print the menu
        display_menu()
        choice=user_choice()
        
        if choice == 1:
            create_doc()
        elif choice == 2:
            add_picture()
        elif choice == 3:
            get_text_doc()

        elif choice == 4:
            print('Exiting the program. Goodbye!')
            break


#--------------------------- Main --------------------------------

menu()

